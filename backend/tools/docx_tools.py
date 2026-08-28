"""Real Word-document generation tool, sandboxed to backend/workspace/.

Produces actual .docx files via python-docx so the agent's deliverables are
openable documents, not chat text or fake extensions.

Expected document structure (see docx_generate):
    Title   -> Heading 1
    Content -> body paragraphs, split on blank lines ("\n\n")

Future extension points (not implemented yet):
    - bullet/numbered lists from a dedicated "bullets" parameter
    - tables from a "table_data" parameter
    - images from a "image_path" parameter
The content-processing block is isolated so those additions only need to
extend the build loop, not restructure this function.
"""

import re
import sys
from pathlib import Path

# Make `tools` importable both when this module is imported by the app
# (backend on sys.path) and when run directly (backend/tools on sys.path).
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from tools.file_tools import WORKSPACE_DIR, _resolve_safe  # noqa: E402

try:
    from docx import Document
except ImportError:
    # python-docx missing: docx_generate reports a clean error at call
    # time instead of failing the whole module import.
    Document = None


def _slugify(title: str) -> str:
    """Turn a title into a safe default filename: lowercase, underscores."""
    slug = title.lower().strip()
    slug = re.sub(r"[^a-z0-9]+", "_", slug)
    slug = slug.strip("_")
    return slug or "document"


def docx_generate(
    title: str,
    content: str,
    filename: str = None,
    provenance_record: dict = None,
) -> dict:
    """Create a Word document in the workspace and save it.

    Args:
        title: Document title, rendered as a Heading 1.
        content: Body text. Split on blank lines ("\n\n") into separate
            paragraphs.
        filename: Optional output name. Defaults to a slug of the title
            (e.g. "Approval Note" -> "approval_note.docx"). ".docx" is
            appended if missing.
        provenance_record: Optional provenance metadata dict. When provided,
            adds a "Provenance & Audit Trail" section at the end of the document.

    Returns:
        {"status": "success", "output": "Document saved as <name>",
         "file_path": "<full path>"} on success, or
        {"status": "error", "output": "<clear error message>"} on failure.
        Never raises.
    """
    if Document is None:
        return {
            "status": "error",
            "output": "python-docx is not installed. Run: pip install python-docx",
        }

    if not content or not content.strip():
        return {
            "status": "error",
            "output": "Content is empty - nothing to put in the document.",
        }

    if filename is None:
        filename = _slugify(title)
    elif not isinstance(filename, str) or not filename.strip():
        filename = _slugify(title)

    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    # Reuse file_tools' sandbox: same workspace root, same traversal
    # rejection, so the deliverable lands next to the agent's other files.
    safe_path = _resolve_safe(filename)
    if safe_path is None:
        return {
            "status": "error",
            "output": f"Path rejected (must stay inside the workspace): {filename!r}",
        }

    try:
        doc = Document()
        doc.add_heading(title, level=1)

        # Split on blank lines so multi-paragraph content renders as real
        # paragraphs instead of one giant block. Single newlines within a
        # paragraph are preserved as line breaks.
        for paragraph_text in content.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue
            doc.add_paragraph(paragraph_text)

        # Append Provenance & Audit Trail section if requested or available
        try:
            from provenance.provenance_tracker import (
                format_provenance_footer,
                get_current_provenance_record,
            )
            if provenance_record is None and get_current_provenance_record is not None:
                provenance_record = get_current_provenance_record()
            
            if provenance_record and format_provenance_footer is not None:
                doc.add_heading("Provenance & Audit Trail", level=3)
                footer_text = format_provenance_footer(provenance_record)
                p = doc.add_paragraph(footer_text)
                try:
                    from docx.shared import Pt, RGBColor
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(128, 128, 128)
                except Exception:
                    pass
        except Exception as prov_exc:
            print(f"[DOCX] Provenance footer skipped: {prov_exc}")

        safe_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(safe_path))
    except OSError as exc:
        return {"status": "error", "output": f"Could not save document: {exc}"}
    except Exception as exc:  # noqa: BLE001 - library errors must not crash.
        return {"status": "error", "output": f"docx_generate failed: {exc}"}

    return {
        "status": "success",
        "output": f"Document saved as {filename}",
        "file_path": str(safe_path),
    }


if __name__ == "__main__":
    TITLE = "Pipe Inspection Approval Note"
    CONTENT = (
        "Findings: Minor corrosion observed on joint A-12.\n\n"
        "Recommendation: Schedule maintenance within 30 days. "
        "No immediate safety risk identified."
    )

    print(f"Workspace: {WORKSPACE_DIR}\n")

    result = docx_generate(TITLE, CONTENT)
    print("RESULT:", result)

    if result["status"] == "success":
        from pathlib import Path

        path = Path(result["file_path"])
        exists = path.exists()
        size = path.stat().st_size if exists else 0
        print(f"\nFile exists: {exists} | Size: {size} bytes")

        # Verify it's a real docx (a ZIP archive with word/document.xml),
        # not just a renamed text file.
        import zipfile

        try:
            with zipfile.ZipFile(path) as zf:
                names = zf.namelist()
                has_document = "word/document.xml" in names
            print(f"Valid docx (ZIP) container: {has_document}")
        except zipfile.BadZipFile:
            print("INVALID: not a real docx (not a ZIP archive)")
